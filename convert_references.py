"""
Convert author-year citations to numbered references in the docx file.
- Preserves all formatting
- Preserves all content except citation style
- Numbers references alphabetically based on bibliography order
"""

import re
from docx import Document
from copy import deepcopy

# Reference mapping: alphabetical order based on the bibliography
# Each entry's index + 1 is its reference number.
# The 'citation' field is the typical citation form used in the document.
REFERENCES = [
    # 1
    {"author_year": "Arafat & Imam, 2022",
     "biblio_prefix": "Arafat, M. E., & Imam, F."},
    # 2
    {"author_year": "Bae & Pyo, 2020",
     "biblio_prefix": "Bae, Y., & Pyo, S."},
    # 3
    {"author_year": "Chow, 2007",
     "biblio_prefix": "Chow, P."},
    # 4
    {"author_year": "Davalos et al., 2001",
     "biblio_prefix": "Davalos, J. F., Qiao, P., Xu, X. F.,"},
    # 5 (alphabetical comes after #4 because Davalos+Qiao < Davalos+Zipfel)
    {"author_year": "Davalos et al., 1999",
     "biblio_prefix": "Davalos, J. F., Zipfel, M. G.,"},
    # 6
    {"author_year": "Dittenber & GangaRao, 2012",
     "biblio_prefix": "Dittenber, D. B., & GangaRao,"},
    # 7
    {"author_year": "Esmaeili et al., 2023",
     "biblio_prefix": "Esmaeili, M. H.,"},
    # 8
    {"author_year": "Ferdous & Manalo, 2014",
     "biblio_prefix": "Ferdous, W., & Manalo, A. (2014)"},
    # 9
    {"author_year": "Ferdous et al., 2016",
     "biblio_prefix": "Ferdous, W., Manalo, A., Aravinthan, T., & Van Erp, G. (2016)"},
    # 10 - Static behaviour (Ferdous + AlAjarmeh)
    {"author_year": "Ferdous et al., 2021a",
     "biblio_prefix": "Ferdous, W., Manalo, A., AlAjarmeh,"},
    # 11 - Geopolymer (Ferdous + Khennane)
    {"author_year": "Ferdous et al., 2015a",
     "biblio_prefix": "Ferdous, W., Manalo, A., Khennane,"},
    # 12 - Book chapter (Ferdous + Muttashar)
    {"author_year": "Ferdous et al., 2021b",
     "biblio_prefix": "Ferdous, W., Manalo, A., Muttashar,"},
    # 13 - Polymers (Ferdous + Salih)
    {"author_year": "Ferdous et al., 2021c",
     "biblio_prefix": "Ferdous, W., Manalo, A., Salih, C., Yu, P., Abousnina,"},
    # 14 - Narrow-gauge (Ferdous + Van Erp + Ghabraie)
    {"author_year": "Ferdous et al., 2018",
     "biblio_prefix": "Ferdous, W., Manalo, A., Van Erp, G., Aravinthan, T., & Ghabraie,"},
    # 15 - Review (Ferdous + Van Erp + Kaewunruen)
    {"author_year": "Ferdous et al., 2015b",
     "biblio_prefix": "Ferdous, W., Manalo, A., Van Erp, G., Aravinthan, T., Kaewunruen,"},
    # 16
    {"author_year": "Gonzalez-Corominas et al., 2017",
     "biblio_prefix": "Gonzalez-Corominas, A.,"},
    # 17
    {"author_year": "Graebe & Woidasky, 2010",
     "biblio_prefix": "Graebe, G., & Woidasky, J."},
    # 18
    {"author_year": "Hugo, 2015",
     "biblio_prefix": "Hugo, A. M."},
    # 19
    {"author_year": "Jabu et al., 2021",
     "biblio_prefix": "Jabu, M. A., Alugongo, A. A., Maube,"},
    # 20
    {"author_year": "Jabu et al., 2024",
     "biblio_prefix": "Jabu, M. A., Alugongo, A. A., & Nkomo,"},
    # 21
    {"author_year": "Jagadeesh et al., 2022",
     "biblio_prefix": "Jagadeesh, P.,"},
    # 22
    {"author_year": "Jing et al., 2021",
     "biblio_prefix": "Jing, G., Siahkouhi, M.,"},
    # 23
    {"author_year": "Ju et al., 2020",
     "biblio_prefix": "Ju, S., Yoon, J.,"},
    # 24
    {"author_year": "Kaewunruen, 2015",
     "biblio_prefix": "Kaewunruen, S. (2015)"},
    # 25
    {"author_year": "Kaewunruen et al., 2017",
     "biblio_prefix": "Kaewunruen, S., You, R.,"},
    # 26
    {"author_year": "Kim, 2024",
     "biblio_prefix": "Kim, J."},
    # 27
    {"author_year": "Koh & Hwang, 2015",
     "biblio_prefix": "Koh, T., & Hwang, S."},
    # 28
    {"author_year": "Koller, 2009",
     "biblio_prefix": "Koller, G. (2009)"},
    # 29
    {"author_year": "Koller, 2015",
     "biblio_prefix": "Koller, G. (2015)"},
    # 30
    {"author_year": "Kondrashchenko et al., 2019",
     "biblio_prefix": "Kondrashchenko, V. I.,"},
    # 31
    {"author_year": "Lampo et al., 2003",
     "biblio_prefix": "Lampo, R., Nosker,"},
    # 32 - Liu, J. (Comparative)
    {"author_year": "Liu, J. et al., 2021",
     "biblio_prefix": "Liu, J., Chen, R.,"},
    # 33 - Liu, X. (Experimental ballast-sleeper)
    {"author_year": "Liu, X. et al., 2021",
     "biblio_prefix": "Liu, X., Kaewunruen,"},
    # 34
    {"author_year": "Louie, 2013",
     "biblio_prefix": "Louie, M. L."},
    # 35
    {"author_year": "Luomala et al., 2024",
     "biblio_prefix": "Luomala, H.,"},
    # 36
    {"author_year": "Manalo & Aravinthan, 2012",
     "biblio_prefix": "Manalo, A. C., & Aravinthan, T."},
    # 37
    {"author_year": "Manalo et al., 2010",
     "biblio_prefix": "Manalo, A., Aravinthan, T., Karunasena,"},
    # 38
    {"author_year": "McConnell, 2008",
     "biblio_prefix": "McConnell, V. P."},
    # 39
    {"author_year": "Palomo et al., 2007",
     "biblio_prefix": "Palomo, A.,"},
    # 40
    {"author_year": "Pattamaprom et al., 2005",
     "biblio_prefix": "Pattamaprom, C.,"},
    # 41
    {"author_year": "Quik et al., 2021",
     "biblio_prefix": "Quik, J. T. K.,"},
    # 42
    {"author_year": "RISSB, 2020",
     "biblio_prefix": "Rail Industry Safety and Standards Board"},
    # 43
    {"author_year": "Saeedi et al., 2024",
     "biblio_prefix": "Saeedi, A.,"},
    # 44
    {"author_year": "Safari et al., 2025",
     "biblio_prefix": "Safari, F.,"},
    # 45
    {"author_year": "Salih et al., 2022",
     "biblio_prefix": "Salih, C., Manalo, A., Ferdous,"},
    # 46
    {"author_year": "Santulli, 2019",
     "biblio_prefix": "Santulli, C."},
    # 47
    {"author_year": "Selig & Waters, 1994",
     "biblio_prefix": "Selig, E. T., & Waters, J. M."},
    # 48
    {"author_year": "Sengsri et al., 2020",
     "biblio_prefix": "Sengsri, P.,"},
    # 49
    {"author_year": "Sharma et al., 2017",
     "biblio_prefix": "Sharma, R. C.,"},
    # 50 - Siahkouhi (carbon fabric, 2022)
    {"author_year": "Siahkouhi et al., 2022",
     "biblio_prefix": "Siahkouhi, M., Li, X., Han,"},
    # 51 - Siahkouhi (KLP, 2021)
    {"author_year": "Siahkouhi et al., 2021",
     "biblio_prefix": "Siahkouhi, M., Li, X., Markine,"},
    # 52
    {"author_year": "Silva et al., 2017",
     "biblio_prefix": "Silva, É. A.,"},
    # 53
    {"author_year": "Singh & Seth, 2023",
     "biblio_prefix": "Singh, A., & Seth, A. K."},
    # 54
    {"author_year": "Suresh Kumar & Muthukannan, 2019",
     "biblio_prefix": "Suresh Kumar, A.,"},
    # 55
    {"author_year": "Thompson et al., 2022",
     "biblio_prefix": "Thompson, S.,"},
    # 56
    {"author_year": "Uehara, 2010",
     "biblio_prefix": "Uehara, M."},
    # 57
    {"author_year": "Van Erp & McKay, 2013",
     "biblio_prefix": "Van Erp, G., & McKay, M."},
    # 58
    {"author_year": "Yu et al., 2021",
     "biblio_prefix": "Yu, P., Manalo,"},
    # 59
    {"author_year": "Zhang et al., 2024",
     "biblio_prefix": "Zhang, D.,"},
]

# Build a simple lookup table for citation strings -> number
# Returns int or None for unmappable
def map_citation_to_number(cite_str, context_para_idx=None):
    """Map an in-text citation string to a reference number.

    Returns int (1-indexed) or None if unmappable.
    """
    cite_str = cite_str.strip()
    # Strip trailing comma/period
    cite_str = cite_str.rstrip('.,;')

    # Direct mapping table (most common)
    direct_map = {
        "Arafat & Imam, 2022": 1,
        "Bae & Pyo, 2020": 2,
        "Chow, 2007": 3,
        "Davalos et al., 2001": 4,
        "Davalos et al., 1999": 5,
        "Dittenber & GangaRao, 2012": 6,
        "Esmaeili et al., 2023": 7,
        "Ferdous & Manalo, 2014": 8,
        "Ferdous et al., 2016": 9,
        "Ferdous et al., 2018": 14,
        "Gonzalez-Corominas et al., 2017": 16,
        "Graebe & Woidasky, 2010": 17,
        "Graebe et al., 2010": 17,  # mis-cited but only one Graebe
        "Hugo, 2015": 18,
        "Jabu et al., 2021": 19,
        "Jabu et al., 2024": 20,
        "Jagadeesh et al., 2022": 21,
        "Jing et al., 2021": 22,
        "Ju et al., 2020": 23,
        "Kaewunruen, 2015": 24,
        "Kaewunruen et al., 2017": 25,
        "Kim, 2024": 26,
        "Koh & Hwang, 2015": 27,
        "Koller, 2009": 28,
        "Koller, 2015": 29,
        "Koller 2015": 29,  # original missing comma
        "Kondrashchenko et al., 2019": 30,
        "Lampo et al., 2003": 31,
        "Louie, 2013": 34,
        "Luomala et al., 2024": 35,
        "Manalo & Aravinthan, 2012": 36,
        "Manalo et al., 2012": 36,  # mis-cited but only Manalo & Aravinthan 2012
        "Manalo et al., 2010": 37,
        "McConnell, 2008": 38,
        "Palomo et al., 2007": 39,
        "Pattamaprom et al., 2005": 40,
        "Quik et al., 2021": 41,
        "RISSB, 2020": 42,
        "Saeedi et al., 2024": 43,
        "Safari et al., 2025": 44,
        "Salih et al., 2022": 45,
        "Santulli, 2019": 46,
        "Selig & Waters, 1994": 47,
        "Selig & Waters 1994": 47,  # original missing comma
        "Sengsri et al., 2020": 48,
        "Sharma et al., 2017": 49,
        "Siahkouhi et al., 2022": 50,
        "Siahkouhi et al., 2021": 51,
        "Silva et al., 2017": 52,
        "Singh & Seth, 2023": 53,
        "Suresh Kumar & Muthukannan, 2019": 54,
        "Thompson et al., 2022": 55,
        "Uehara, 2010": 56,
        "Van Erp & McKay, 2013": 57,
        "Yu et al., 2021": 58,
        "Zhang et al., 2024": 59,
    }

    if cite_str in direct_map:
        return direct_map[cite_str]

    # Ambiguous: Ferdous et al., 2015 -> #15 (review) or #11 (geopolymer)
    # Default to review (#15). Use #11 for explicit geopolymer figure caption.
    if cite_str == "Ferdous et al., 2015":
        if context_para_idx == 113:
            return 11  # figure caption (c) Geopolymer specifically
        return 15  # default = review paper

    # Ambiguous: Ferdous et al., 2021 -> #10, #12, or #13
    # Default to book chapter on alternatives (#12) - most general
    if cite_str == "Ferdous et al., 2021":
        return 12

    # Ambiguous: Liu et al., 2021 -> #32 (Liu,J) or #33 (Liu,X)
    if cite_str == "Liu et al., 2021":
        if context_para_idx == 138:
            return 32  # lateral resistance - Liu, J
        return 33  # ballast interaction - Liu, X (default)

    # Best-effort mapping for likely typos / variants
    if cite_str == "Ferdous et al., 2020":
        return 12  # likely typo for 2021 - book chapter
    if cite_str == "Salih et al., 2021":
        return 13  # Salih is co-author of Ferdous et al., 2021 Polymers
    if cite_str == "Siahkouhi et al., 2026":
        return 51  # likely typo for 2021 - KLP review

    # Unmappable
    return None


# Citations that are NOT in the bibliography at all - keep as author-year
UNMAPPABLE = {
    "Kaewunruen et al., 2024",
    "Giunta & Rouisse, 2026",
}


def parse_citation_group(text_inside_parens, context_para_idx=None):
    """Parse the inside of a citation group (between parens or brackets).

    Returns: (prefix_text, list_of_numbers, list_of_unmapped, suffix_text)
        prefix_text: any text BEFORE the first author-year (e.g., "data extracted from ")
        list_of_numbers: numbers for mappable citations
        list_of_unmapped: list of (position, original_text) for unmappable
        suffix_text: any text AFTER (typically empty)
    """
    # Split by ; or ,(uncommon between separate cites). Use ; as primary separator.
    # But careful: "Author1, Author2 et al., 2020; Author3, 2021"
    # Heuristic: split on '; '
    raw_parts = [p.strip() for p in re.split(r'\s*;\s*', text_inside_parens) if p.strip()]

    numbers = []
    unmapped = []
    prefix = ""

    for part_idx, part in enumerate(raw_parts):
        # Check for prefix like "data extracted from"
        # Pattern: "data extracted from Author et al., Year"
        m = re.match(r'^(data extracted from\s+)(.+)$', part, re.IGNORECASE)
        if m:
            prefix = m.group(1)
            part = m.group(2).strip()

        # part should be like "Author et al., 2020" or "Author & Author, 2020" or "Author, 2020"
        num = map_citation_to_number(part, context_para_idx)
        if num is not None:
            numbers.append(num)
        elif part in UNMAPPABLE:
            unmapped.append(part)
        else:
            # Unknown - try to gracefully keep
            unmapped.append(part)

    return prefix, numbers, unmapped


def build_replacement(text_inside_brackets, opening_char, closing_char, context_para_idx=None):
    """Build the replacement string for an entire citation group.

    text_inside_brackets: text between the opening and closing chars (not including them)
    opening_char: '(' or '['
    closing_char: ')' or ']'

    Returns the full replacement string (including brackets if needed).
    """
    prefix, numbers, unmapped = parse_citation_group(text_inside_brackets, context_para_idx)

    # Sort numbers ascending and deduplicate
    numbers = sorted(set(numbers))

    # If we have a prefix or unmapped, preserve the original structure
    if prefix:
        # e.g. "(data extracted from [51])"
        if numbers and not unmapped:
            return f"{opening_char}{prefix}[{', '.join(str(n) for n in numbers)}]{closing_char}"
        elif numbers and unmapped:
            inner = ', '.join(str(n) for n in numbers)
            unmapped_str = '; '.join(unmapped)
            return f"{opening_char}{prefix}[{inner}]; {unmapped_str}{closing_char}"
        elif unmapped:
            unmapped_str = '; '.join(unmapped)
            return f"{opening_char}{prefix}{unmapped_str}{closing_char}"
        else:
            return f"{opening_char}{text_inside_brackets}{closing_char}"

    # No prefix - simple case
    if numbers and not unmapped:
        # Standard case: "[N1, N2, ...]" - no need for outer parens since brackets are the citation
        return f"[{', '.join(str(n) for n in numbers)}]"
    elif numbers and unmapped:
        # Mixed: numbers + unmapped
        inner_nums = ', '.join(str(n) for n in numbers)
        unmapped_str = '; '.join(unmapped)
        return f"[{inner_nums}]; ({unmapped_str})"
    elif unmapped:
        # Only unmapped - keep original style
        unmapped_str = '; '.join(unmapped)
        return f"{opening_char}{unmapped_str}{closing_char}"
    else:
        # No matches at all - return original
        return f"{opening_char}{text_inside_brackets}{closing_char}"


def is_citation_group(content):
    """Check if the content between brackets looks like a citation group."""
    # Must contain at least one year (4-digit number, possibly with a/b/c)
    if not re.search(r'\b(19|20)\d{2}[a-z]?\b', content):
        return False
    # Must contain at least one author-like word (capital letter start)
    if not re.search(r'[A-Z][a-z]+', content):
        return False
    return True


def replace_citations_in_text(text, context_para_idx=None):
    """Replace all citation groups in a text string.

    Returns the modified text.
    """
    # Process round-bracketed and square-bracketed citations
    # We need to be careful to only replace citation-like groups, not other parens

    # First, find all candidate groups
    def replacer(m):
        opening = m.group(1)
        content = m.group(2)
        closing = m.group(3)
        if not is_citation_group(content):
            return m.group(0)
        return build_replacement(content, opening, closing, context_para_idx)

    # Process parens
    text = re.sub(r'(\()([^()]*?)(\))', replacer, text)
    # Process square brackets
    text = re.sub(r'(\[)([^\[\]]*?)(\])', replacer, text)

    return text


def replace_citations_in_paragraph(para, para_idx=None):
    """Apply citation replacements to a paragraph (handling multi-run citations).

    para: a python-docx Paragraph object
    para_idx: optional paragraph index for context-aware mapping
    """
    if not para.runs:
        return

    # Get full text and character-to-run mapping
    char_to_run = []
    full_text = ""
    for run_idx, run in enumerate(para.runs):
        for pos in range(len(run.text)):
            char_to_run.append((run_idx, pos))
        full_text += run.text

    # Compute new full text using text-level replacement
    new_full_text = replace_citations_in_text(full_text, para_idx)

    if new_full_text == full_text:
        return  # nothing to do

    # Find all replacement regions
    # We re-find them by comparing the two strings
    # Simpler approach: find all citation groups in original, compute their replacements,
    # then apply each replacement to the runs.

    # Collect (start_in_full, end_in_full, replacement_str) for each citation group
    replacements = []

    def collector(m, brackets):
        start = m.start()
        end = m.end()
        opening = m.group(1)
        content = m.group(2)
        closing = m.group(3)
        if not is_citation_group(content):
            return
        replacement = build_replacement(content, opening, closing, para_idx)
        replacements.append((start, end, replacement))

    for m in re.finditer(r'(\()([^()]*?)(\))', full_text):
        collector(m, 'paren')
    for m in re.finditer(r'(\[)([^\[\]]*?)(\])', full_text):
        collector(m, 'square')

    if not replacements:
        return

    # Sort by start descending so we replace from the end
    replacements.sort(key=lambda x: x[0], reverse=True)

    for start, end, replacement in replacements:
        first_run_idx, first_pos = char_to_run[start]
        last_run_idx, last_pos = char_to_run[end - 1]

        if first_run_idx == last_run_idx:
            # Single-run replacement
            run = para.runs[first_run_idx]
            run.text = run.text[:first_pos] + replacement + run.text[last_pos + 1:]
        else:
            # Multi-run replacement
            first_run = para.runs[first_run_idx]
            last_run = para.runs[last_run_idx]
            text_before = first_run.text[:first_pos]
            text_after = last_run.text[last_pos + 1:]
            first_run.text = text_before + replacement
            for idx in range(first_run_idx + 1, last_run_idx):
                para.runs[idx].text = ""
            last_run.text = text_after


def is_bibliography_entry(para):
    """Check if a paragraph is a bibliography entry (List Paragraph style with author start)."""
    if para.style.name != "List Paragraph":
        return False
    text = para.text.strip()
    # Must contain a year pattern like "(YYYY)"
    if not re.search(r'\((?:19|20)\d{2}\)', text):
        return False
    return True


def number_bibliography(doc):
    """Add reference numbers to bibliography entries based on alphabetical mapping."""
    # Build a list of bibliography entry paragraphs
    biblio_paras = []
    for i, para in enumerate(doc.paragraphs):
        if is_bibliography_entry(para):
            biblio_paras.append((i, para))

    # For each entry, find which reference number it corresponds to
    # by matching the prefix
    for para_idx, para in biblio_paras:
        text = para.text.strip()
        matched_ref = None
        matched_idx = None
        for i, ref in enumerate(REFERENCES):
            if text.startswith(ref["biblio_prefix"]):
                if matched_ref is None or len(ref["biblio_prefix"]) > len(matched_ref["biblio_prefix"]):
                    matched_ref = ref
                    matched_idx = i

        if matched_ref is None:
            print(f"WARNING: Could not match bibliography entry at para {para_idx}: {text[:120]}")
            continue

        ref_number = matched_idx + 1

        # Add [N] prefix to the first run that has non-empty text
        prefix_str = f"[{ref_number}] "
        for run in para.runs:
            if run.text:
                run.text = prefix_str + run.text
                break


def main():
    doc = Document('original.docx')

    # Process each body paragraph
    for i, para in enumerate(doc.paragraphs):
        replace_citations_in_paragraph(para, para_idx=i)

    # Process each table cell
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_citations_in_paragraph(para)

    # Add numbers to bibliography
    number_bibliography(doc)

    doc.save('numbered.docx')
    print("Saved: numbered.docx")


if __name__ == "__main__":
    main()
